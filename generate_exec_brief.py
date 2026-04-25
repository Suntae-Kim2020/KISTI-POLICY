#!/usr/bin/env python3
"""
KISTI 정책본부장용 시연 자료 생성 스크립트
산출물: KISTI_Policy_대시보드_정책본부장보고.docx
"""
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ACCENT = RGBColor(0x1F, 0x4E, 0x79)   # 진한 파랑 (KISTI 브랜드 톤)
MUTED = RGBColor(0x59, 0x59, 0x59)
HIGHLIGHT = RGBColor(0xC0, 0x39, 0x2B)  # 강조 수치 (적색)


def set_cell_bg(cell, rgb_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_hex)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT
    return h


def add_p(doc, text, bold=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r1 = p.add_run(bold_lead)
        r1.bold = True
        r1.font.color.rgb = ACCENT
        p.add_run(" — " + text)
    else:
        p.add_run(text)
    return p


def add_table(doc, headers, rows, widths_cm=None, header_color="1F4E79"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Light Grid Accent 1"

    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        set_cell_bg(hdr[i], header_color)
        if widths_cm and i < len(widths_cm):
            hdr[i].width = Cm(widths_cm[i])

    for r_idx, row in enumerate(rows):
        cells = t.rows[r_idx + 1].cells
        for i, v in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            run = p.add_run(str(v))
            run.font.size = Pt(10)
            if widths_cm and i < len(widths_cm):
                cells[i].width = Cm(widths_cm[i])
    return t


def add_callout(doc, text, color_hex="FFF4E1"):
    """강조 박스 (연한 배경 + 진한 글자)"""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, color_hex)
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = ACCENT
    return t


def main():
    doc = Document()

    # 페이지 여백
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ── 표지/제목 ─────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("KISTI 인프라 정책 가치 분석 대시보드")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = ACCENT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("— 정책본부장 시연 보고 자료 —")
    r.font.size = Pt(13)
    r.font.color.rgb = MUTED

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("서비스 URL: https://kisti.ailibrary.kr\n작성: 2026년 4월")
    r.font.size = Pt(10)
    r.font.color.rgb = MUTED

    doc.add_paragraph()

    # ── 1. Executive Summary ──────────────────────────────
    add_heading(doc, "1. 한눈에 보기", level=1)
    add_callout(
        doc,
        "“KISTI 인프라(슈퍼컴·연구망·e-Science 플랫폼)가 유발한 연구 성과를 "
        "정량 지표로 입증하고, 국내외 유사 기관과 비교해 "
        "정책적 투자 가치를 데이터로 설득하는 분석 서비스.”",
    )
    doc.add_paragraph()

    add_p(doc, "■ 핵심 수치 (2008–2024 기준)", bold=True, size=12)
    add_table(
        doc,
        headers=["지표", "값", "정책적 함의"],
        rows=[
            ["KISTI 유발논문 수", "6,104편", "직접논문(1,956편)의 3.1배 — 숨은 기여 드러냄"],
            ["MNCS (피인용 영향력)", "1.957", "한국 평균 대비 95.7% 높은 영향력"],
            ["논문 생산성 / 10억원·년", "9.43편", "인프라 예산(KSC+KREONET 381억원) 기준"],
            ["피인용 생산성 / 10억원·년", "291.44회", "예산당 지식 파급 효과"],
            ["국제 비교 (XSEDE, 2011–2022)", "MNCS 1.975", "미국 XSEDE와 동등한 질적 수준"],
        ],
        widths_cm=[5.0, 3.5, 7.5],
    )

    doc.add_paragraph()

    # ── 2. 문제의식 ──────────────────────────────────────
    add_heading(doc, "2. 왜 이 분석이 필요한가", level=1)
    add_p(doc, "KISTI 인프라의 가치는 전통적 지표로는 보이지 않습니다.", bold=True)
    add_bullet(
        doc,
        "슈퍼컴퓨터·연구망은 직접 논문을 쓰지 않는다. 타 연구자들이 쓴다.",
        bold_lead="구조적 한계",
    )
    add_bullet(
        doc,
        "KISTI 소속 저자 논문 수(직접논문)만 집계하면 실제 기여의 24%만 드러난다.",
        bold_lead="가시성 문제",
    )
    add_bullet(
        doc,
        "예산 대비 성과를 정량화할 공식 지표가 부재해 매년 예산 방어가 서사 중심으로 흐른다.",
        bold_lead="정책 설득력",
    )

    add_p(doc, "■ 본 프로젝트의 해답", bold=True, size=12)
    add_bullet(
        doc,
        "논문의 사사 표기(Funding Text: FU/FX)에서 KISTI 인프라 사용 키워드를 자동 검출",
        bold_lead="유발논문 개념",
    )
    add_bullet(
        doc,
        "한국 전체 논문의 분야·연도별 평균 피인용 대비 영향력을 계산 (MNCS)",
        bold_lead="영향력 정규화",
    )
    add_bullet(
        doc,
        "예산 규모 차이 보정 (논문·피인용 / 10억원·년)",
        bold_lead="투자 효율 지표",
    )
    add_bullet(
        doc,
        "XSEDE, PRACE, NERSC 등 국외 주요 인프라와 동일 기간·동일 기준 비교",
        bold_lead="국제 벤치마킹",
    )

    doc.add_paragraph()

    # ── 3. 분석 프레임 ───────────────────────────────────
    add_heading(doc, "3. 분석 프레임", level=1)

    add_p(doc, "■ 4-way 비교 구조", bold=True, size=12)
    add_table(
        doc,
        headers=["기관", "특성", "분석 섹션", "비교 의의"],
        rows=[
            ["KISTI", "정보 인프라 (KSC, KREONET, EDISON)", "sec1~sec3 (11페이지)", "주 대상"],
            ["KBSI", "분석 장비 서비스", "sec4~sec6 (15페이지)", "유사 성격 인프라"],
            ["IBS", "독립 연구소 (자체 논문)", "sec7~sec9 (14페이지)", "인프라 vs 연구소 대비"],
            ["PAL", "가속기 (POSTECH 하위)", "sec10 (6페이지)", "또 다른 인프라 유형"],
        ],
        widths_cm=[2.5, 5.0, 4.0, 4.5],
    )

    add_p(doc, "■ 데이터 규모", bold=True, size=12)
    add_bullet(doc, "WoS SCIE 한국 논문 955,146건 (2008~2024, 17년치)")
    add_bullet(doc, "기관별 매핑 281만건, ESI 22분야 분류")
    add_bullet(doc, "JCR JIF/Quartile, HCP, 국제공동연구 등 다각도 지표")

    add_p(doc, "■ 분류 원칙 (엄격함)", bold=True, size=12)
    add_bullet(
        doc,
        "KISTI 소속 저자가 있어도 사사표기에 인프라 키워드가 있으면 유발논문으로 분류",
        bold_lead="중복 방지",
    )
    add_bullet(
        doc,
        "직접논문과 유발논문이 서로 중복 집계되지 않도록 624편 제외 처리",
    )

    doc.add_paragraph()

    # ── 4. 차별화 가치 ───────────────────────────────────
    add_heading(doc, "4. 정책적 차별화 가치", level=1)

    add_p(doc, "① 숨겨진 기여의 정량화", bold=True, size=12)
    add_p(
        doc,
        "KISTI 소속 논문 1,956편만 보면 기여 과소 평가. 유발논문 6,104편을 더하면 "
        "실제 기여는 8,060편. 인프라 투자 효과가 3배 가까이 더 크게 드러납니다.",
    )

    add_p(doc, "② 질적 우수성 검증", bold=True, size=12)
    add_p(
        doc,
        "MNCS 1.957은 한국 평균 대비 95.7% 높은 피인용 영향력. "
        "즉 KISTI 인프라로 지원된 연구가 한국 평균보다 훨씬 높은 학술적 인용을 받고 있음을 의미. "
        "Q1 저널 비중도 유발논문 64.0% (직접논문 53.4%)로 더 높음.",
    )

    add_p(doc, "③ 예산 대비 효율 입증", bold=True, size=12)
    add_p(
        doc,
        "KISTI 인프라 연간 예산 약 381억원(KSC 300억 + KREONET 81억) 기준, "
        "연 9.43편의 고품질 논문이 유발됨. XSEDE·PRACE 등 글로벌 벤치마크와 "
        "10억원당 정규화로 직접 비교 가능.",
    )

    add_p(doc, "④ 국제 경쟁력 데이터", bold=True, size=12)
    add_table(
        doc,
        headers=["프로그램", "기간", "KISTI MNCS", "상대 비교 의의"],
        rows=[
            ["XSEDE (미국)", "2011–2022", "1.975", "미국 최대 학술용 슈퍼컴 네트워크와 동등 수준"],
            ["PRACE (유럽)", "2010–2023", "1.987", "유럽 Tier-0 센터와 대등한 영향력"],
            ["HECToR (영국)", "2008–2014", "1.365", "참고용 (해당 기간 SCIE 원시데이터 부재)"],
        ],
        widths_cm=[3.5, 3.0, 2.5, 7.0],
    )

    doc.add_paragraph()

    # ── 5. 분석 페이지 구성 (52개) ────────────────────────
    add_heading(doc, "5. 분석 대시보드 구성 (52페이지)", level=1)
    add_table(
        doc,
        headers=["섹션", "제목", "주요 내용"],
        rows=[
            ["요약", "요약 대시보드", "KISTI/KBSI/IBS 3개 기관 요약 카드 + 연도별 합산 차트"],
            ["1", "KISTI 논문분석 (5)", "발표 현황, 분야별, 영향력, 협력, 학술지"],
            ["2", "KISTI 유발논문분석 (8)", "인프라별, 수혜기관, MNCS, 예산 정규화, 국제 비교"],
            ["3", "KISTI 비교·종합 (3)", "직접 vs 유발, 기여도, 인프라 투자 효과"],
            ["4–5", "KBSI 직접·유발 (11)", "KISTI와 동일 분석 프레임으로 대조"],
            ["6", "KISTI vs KBSI 비교 (4)", "인력 생산성 포함"],
            ["7–8", "IBS 직접·유발 (11)", "연구소 형태 기관과의 대비"],
            ["9", "KISTI vs IBS 비교 (3)", "인프라형 vs 연구소형 구조 비교"],
            ["10", "PAL 유발논문분석 (6)", "가속기 인프라 유형 비교"],
            ["11", "질적 우수성 종합", "7개 기관·유형 종합 비교 테이블"],
        ],
        widths_cm=[2.2, 5.8, 8.0],
    )

    doc.add_paragraph()

    # ── 6. 시스템 구성 및 보안 ───────────────────────────
    add_heading(doc, "6. 시스템 운영 현황", level=1)

    add_p(doc, "■ 운영 인프라", bold=True, size=12)
    add_bullet(doc, "https://kisti.ailibrary.kr", bold_lead="프로덕션")
    add_bullet(doc, "https://dev.kisti.ailibrary.kr (내부 검증용)", bold_lead="개발/스테이징")
    add_bullet(doc, "Google Cloud Platform (Cloud Run, 서울→도쿄 리전)", bold_lead="호스팅")

    add_p(doc, "■ 접근 통제", bold=True, size=12)
    add_bullet(doc, "이메일/비밀번호 로그인 필수. 로그인 실패 5회 시 10분 자동 잠금")
    add_bullet(doc, "관리자·일반 사용자 2단계 권한 분리")
    add_bullet(doc, "로그인·다운로드 등 모든 활동 감사 로그 기록 (접속 IP, 시각 포함)")
    add_bullet(doc, "브라우저 세션 HTTPS·HttpOnly·SameSite 적용")

    add_p(doc, "■ 데이터 갱신 체계", bold=True, size=12)
    add_bullet(doc, "매년 KISTEP 본 과제와 연동되어 신규 WoS 데이터 갱신")
    add_bullet(doc, "한 번의 명령(rebuild)으로 원본→통계→배포까지 자동화")
    add_bullet(doc, "개발환경(dev) 검증 후 프로덕션 배포 — 서비스 중단 없음")

    doc.add_paragraph()

    # ── 7. 시연 시나리오 ─────────────────────────────────
    add_heading(doc, "7. 시연 시나리오 (약 10분)", level=1)

    add_table(
        doc,
        headers=["시간", "화면", "핵심 메시지"],
        rows=[
            ["0:00–1:00", "로그인 + 요약 대시보드", "KISTI/KBSI/IBS 3기관 요약 카드 한눈에 비교"],
            ["1:00–3:00", "sec2_7 경제적 가치 추정", "MNCS 1.957 + 예산 정규화 9.43편/10억원"],
            ["3:00–5:00", "sec2_8 국제 비교", "XSEDE·PRACE와 기간 매칭 비교 — 국제 경쟁력 입증"],
            ["5:00–7:00", "sec11 질적 우수성 종합", "4개 기관 7개 유형 전체 비교 테이블"],
            ["7:00–8:30", "라이브차트 HTML 다운로드", "서버 없이 공유 가능한 독립 보고서 생성"],
            ["8:30–10:00", "사용자 관리 + 감사 로그", "내부 공유·보안 통제 수준 시연"],
        ],
        widths_cm=[2.2, 4.8, 9.0],
    )

    doc.add_paragraph()

    # ── 8. 핵심 메시지 정리 ──────────────────────────────
    add_heading(doc, "8. 정책본부장께 드리는 핵심 메시지", level=1)

    add_callout(
        doc,
        "① KISTI 인프라의 실제 기여는 공식 지표보다 3배 크다 "
        "(직접논문 1,956편 + 유발논문 6,104편 = 실제 기여 8,060편).",
    )
    doc.add_paragraph()
    add_callout(
        doc,
        "② 유발논문은 한국 평균보다 95.7% 더 많이 인용된다 (MNCS 1.957). "
        "즉 KISTI 인프라는 ‘많이’가 아니라 ‘잘’ 지원한다.",
    )
    doc.add_paragraph()
    add_callout(
        doc,
        "③ 투자 효율은 XSEDE·PRACE 등 해외 동급 인프라와 대등하다. "
        "매년 381억 투자 → 58편/년, 피인용 17,930건/년 발생.",
    )
    doc.add_paragraph()
    add_callout(
        doc,
        "④ 이 지표들을 매년 갱신하여 예산 협의·정책 보고의 공식 근거로 활용 가능하다. "
        "KISTEP 본 과제와 자동 연동되어 운영 부담 최소화.",
    )

    doc.add_paragraph()

    # ── 9. 향후 확장 ──────────────────────────────────────
    add_heading(doc, "9. 확장 가능성", level=1)
    add_bullet(doc, "NASA, RIKEN, AIST 등 추가 벤치마크 기관 연동", bold_lead="비교군 확장")
    add_bullet(doc, "연구자 단위 분석 (현재 기관 단위). WoS 데이터 필드 확장 필요", bold_lead="세분화")
    add_bullet(doc, "Scopus, 특허, 사업화 실적 등 데이터 출처 다양화", bold_lead="다원화")
    add_bullet(doc, "정기 자동 리포트 이메일 발송 (월간/분기)", bold_lead="자동 보고")

    doc.add_paragraph()

    # ── 10. 부록 ──────────────────────────────────────────
    add_heading(doc, "10. 부록 — 주요 지표 설명", level=1)

    add_p(doc, "■ MNCS (Mean Normalized Citation Score)", bold=True, size=11)
    add_p(
        doc,
        "각 논문의 피인용수를 “같은 분야·같은 연도 평균 피인용”으로 정규화한 뒤 평균. "
        "분야·연도 편차를 제거해 순수 학술적 영향력만 측정. "
        "1.0 = 평균, 2.0 = 평균의 2배 수준. "
        "국제 표준 지표(CNCI, FWCI)와 수학적으로 동등.",
        size=10,
    )

    add_p(doc, "■ 유발논문 정의", bold=True, size=11)
    add_p(
        doc,
        "WoS의 사사 표기 필드(FU/FX)에서 KISTI 인프라 키워드(KSC, NURION, KREONET, EDISON 등)가 "
        "검출된 모든 논문. KISTI 소속 공저자 유무와 무관. "
        "실질적 인프라 사용에 기반한 객관적 판별 방식.",
        size=10,
    )

    add_p(doc, "■ ESI 22분야 분류", bold=True, size=11)
    add_p(
        doc,
        "Clarivate ESI 기준 22개 연구 분야 (Physics, Chemistry, Engineering, Clinical Medicine 등). "
        "분야 간 피인용 문화 차이를 보정하는 기준.",
        size=10,
    )

    add_p(doc, "■ Q1 저널 비율", bold=True, size=11)
    add_p(
        doc,
        "JCR Journal Impact Factor 상위 25% 저널 게재 논문의 비율. "
        "학술지 수준 기반 질적 지표.",
        size=10,
    )

    doc.add_paragraph()
    doc.add_paragraph()

    # ── 문의 ──────────────────────────────────────────────
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer_p.add_run(
        "본 서비스·자료 관련 문의: 김선태 (전북대학교)  ·  kim.suntae@jbnu.ac.kr"
    )
    r.font.size = Pt(10)
    r.font.color.rgb = MUTED

    # 저장
    out_path = "/Users/kimsuntae/KISTI_Policy/KISTI_Policy_대시보드_정책본부장보고.docx"
    doc.save(out_path)
    print(f"생성 완료: {out_path}")


if __name__ == "__main__":
    main()
